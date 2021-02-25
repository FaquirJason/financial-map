import re, jieba
import jieba.posseg as pseg
from pyltp import SentenceSplitter

import pymongo
from sentence_parser import *

import os, re
from ddparser import DDParser

from docx import Document
import xlwt

from multiprocessing import Pool

class CausalityExractor():
    def __init__(self):
        pass

    '''1由果溯因配套式'''
    def ruler1(self, sentence):
        '''
        conm2:〈[之]所以,因为〉、〈[之]所以,由于〉、 <[之]所以,缘于〉
        conm2_model:<Conj>{Effect},<Conj>{Cause}
        '''
        datas = list()
        word_pairs =[['之?所以', '因为'], ['之?所以', '由于'], ['之?所以', '缘于']]
        for word in word_pairs:
            pattern = re.compile(r'\s?(%s)/[p|c]+\s(.*)(%s)/[p|c]+\s(.*)' % (word[0], word[1]))
            result = pattern.findall(sentence)
            data = dict()
            if result:
                data['tag'] = result[0][0] + '-' + result[0][2]
                data['cause'] = result[0][3]
                data['effect'] = result[0][1]
                datas.append(data)
        if datas:
            return datas[0]
        else:
            return {}
    '''2由因到果配套式'''
    def ruler2(self, sentence):
        '''
        conm1:〈因为,从而〉、〈因为,为此〉、〈既[然],所以〉、〈因为,为此〉、〈由于,为此〉、〈只有|除非,才〉、〈由于,以至[于]>、〈既[然],却>、
        〈如果,那么|则〉、<由于,从而〉、<既[然],就〉、〈既[然],因此〉、〈如果,就〉、〈只要,就〉〈因为,所以〉、 <由于,于是〉、〈因为,因此〉、
         <由于,故〉、 〈因为,以致[于]〉、〈因为,因而〉、〈由于,因此〉、<因为,于是〉、〈由于,致使〉、〈因为,致使〉、〈由于,以致[于] >
         〈因为,故〉、〈因[为],以至[于]>,〈由于,所以〉、〈因为,故而〉、〈由于,因而〉
        conm1_model:<Conj>{Cause}, <Conj>{Effect}
        '''
        datas = list()
        word_pairs =[['因为', '从而'], ['因为', '为此'], ['既然?', '所以']
        ,
                    ['因为', '为此'], ['由于', '为此'], ['除非', '才'],
                    ['只有', '才'], ['由于', '以至于?'], ['既然?', '却'],
                    ['如果', '那么'], ['如果', '则'], ['由于', '从而'],
                    ['既然?', '就'], ['既然?', '因此'], ['如果', '就'],
                    ['只要', '就'], ['因为', '所以'], ['由于', '于是'],
                    ['因为', '因此'], ['由于', '故'], ['因为', '以致于?'],
                    ['因为', '以致'], ['因为', '因而'], ['由于', '因此'],
                    ['因为', '于是'], ['由于', '致使'], ['因为', '致使'],
                    ['由于', '以致于?'], ['因为', '故'], ['因为?', '以至于?'],
                    ['由于', '所以'], ['因为', '故而'], ['由于', '因而']]

        for word in word_pairs:
            pattern = re.compile(r'\s?(%s)/[p|c]+\s(.*)(%s)/[p|c]+\s(.*)' % (word[0], word[1]))
            result = pattern.findall(sentence)
            data = dict()
            if result:
                data['tag'] = result[0][0] + '-' + result[0][2]
                data['cause'] = result[0][1]
                data['effect'] = result[0][3]
                datas.append(data)
        if datas:
            return datas[0]
        else:
            return {}
    '''3由因到果居中式明确'''
    def ruler3(self, sentence):
        '''
        cons2:于是、所以、故、致使、以致[于]、因此、以至[于]、从而、因而
        cons2_model:{Cause},<Conj...>{Effect}
        '''

        pattern = re.compile(r'(.*)[,，]+.*(于是|所以|故|致使|以致于?|以至于?|从而)/[p|c]+\s(.*)')
        result = pattern.findall(sentence)
        data = dict()
        if result:
            data['tag'] = result[0][1]
            data['cause'] = result[0][0]
            data['effect'] = result[0][2]
        return data
    '''4由因到果居中式精确'''
    def ruler4(self, sentence):
        '''
        verb1:牵动、导向、使动、导致、勾起、引入、指引、使、予以、产生、促成、造成、引导、造就、促使、酿成、
            引发、渗透、促进、引起、诱导、引来、促发、引致、诱发、推进、诱致、推动、招致、影响、致使、滋生、归于、
            作用、使得、决定、攸关、令人、引出、浸染、带来、挟带、触发、关系、渗入、诱惑、波及、诱使
        verb1_model:{Cause},<Verb|Adverb...>{Effect}
        '''
        pattern = re.compile(r'(.*)\s+(牵动|已致|导向|使动|导致|勾起|引入|指引|使|予以|产生|促成|造成|引导|造就|促使|酿成|引发|渗透|促进|引起|诱导|引来|促发|引致|诱发|推进|诱致|推动|招致|影响|致使|滋生|归于|作用|使得|决定|攸关|令人|引出|浸染|带来|挟带|触发|关系|渗入|诱惑|波及|诱使)/[d|v]+\s(.*)')
        result = pattern.findall(sentence)
        data = dict()
        if result:
            data['tag'] = result[0][1]
            data['cause'] = result[0][0]
            data['effect'] = result[0][2]
        return data
    '''5由因到果前端式模糊'''
    def ruler5(self, sentence):
        '''
        prep:为了、依据、为、按照、因[为]、按、依赖、照、比、凭借、由于
        prep_model:<Prep...>{Cause},{Effect}
        '''
        pattern = re.compile(r'\s?(为了|依据|按照|因为|因|按|依赖|凭借|由于)/[p|c]+\s(.*)[,，]+(.*)')
        result = pattern.findall(sentence)
        data = dict()
        if result:
            data['tag'] = result[0][0]
            data['cause'] = result[0][1]
            data['effect'] = result[0][2]

        return data

    '''6由因到果居中式模糊'''
    def ruler6(self, sentence):
        '''
        adverb:以免、以便、为此、才
        adverb_model:{Cause},<Verb|Adverb...>{Effect}
        '''
        pattern = re.compile(r'(.*)(以免|以便|为此|才)\s(.*)')
        result = pattern.findall(sentence)
        data = dict()
        if result:
            data['tag'] = result[0][1]
            data['cause'] = result[0][0]
            data['effect'] = result[0][2]
        return data

    '''7由因到果前端式精确'''
    def ruler7(self, sentence):
        '''
        cons1:既[然]、因[为]、如果、由于、只要
        cons1_model:<Conj...>{Cause},{Effect}
        '''
        pattern = re.compile(r'\s?(既然?|因|因为|如果|由于|只要)/[p|c]+\s(.*)[,，]+(.*)')
        result = pattern.findall(sentence)
        data = dict()
        if result:
            data['tag'] = result[0][0]
            data['cause'] = result[0][1]
            data['effect'] = result[0][2]
        return data
    '''8由果溯因居中式模糊'''
    def ruler8(self, sentence):
        '''
        3
        verb2:根源于、取决、来源于、出于、取决于、缘于、在于、出自、起源于、来自、发源于、发自、源于、根源于、立足[于]
        verb2_model:{Effect}<Prep...>{Cause}
        '''

        pattern = re.compile(r'(.*)(根源于|取决|来源于|出于|取决于|缘于|在于|出自|起源于|来自|发源于|发自|源于|根源于|立足|立足于)/[p|c]+\s(.*)')
        result = pattern.findall(sentence)
        data = dict()
        if result:
            data['tag'] = result[0][1]
            data['cause'] = result[0][2]
            data['effect'] = result[0][0]
        return data
    '''9由果溯因居端式精确'''
    def ruler9(self, sentence):
        '''
        cons3:因为、由于
        cons3_model:{Effect}<Conj...>{Cause}
        '''
        pattern = re.compile(r'(.*)是?\s(因为|由于)/[p|c]+\s(.*)')
        result = pattern.findall(sentence)
        data = dict()
        if result:
            data['tag'] = result[0][1]
            data['cause'] = result[0][2]
            data['effect'] = result[0][0]

        return data

    '''抽取主函数'''
    def extract_triples(self, sentence):
        infos = list()
      #  print(sentence)
        if self.ruler1(sentence):
            infos.append(self.ruler1(sentence))
        elif self.ruler2(sentence):
            infos.append(self.ruler2(sentence))
        elif self.ruler3(sentence):
            infos.append(self.ruler3(sentence))
        # elif self.ruler4(sentence):
        #     infos.append(self.ruler4(sentence))
        # elif self.ruler5(sentence):
        #     infos.append(self.ruler5(sentence))
        elif self.ruler6(sentence):
            infos.append(self.ruler6(sentence))
        # elif self.ruler7(sentence):
        #     infos.append(self.ruler7(sentence))
        # elif self.ruler8(sentence): 
        # elif self.ruler9(sentence):
        #     infos.append(self.ruler9(sentence))

        return infos

    '''抽取主控函数'''
    def extract_main(self, content):
        sentences = self.process_content(content)
        datas = list()
        
        for sentence in sentences:
            index = sentences.index(sentence)
            subsents = self.fined_sentence(sentence)
            subsents.append(sentence)
            for sent in subsents:
                sent = ' '.join([word.word + '/' + word.flag for word in pseg.cut(sent)])
                result = self.extract_triples(sent)
                if result:
                    for data in result:
                        if data['tag'] and data['cause'] and data['effect']:
                            if index != 0 and index != len(sentences)-1:
                                data["around"] = sentences[index-1]+sentences[index]+sentences[index+1]
                            elif index == 0 :
                                data["around"] = sentences[index]+sentences[index+1]
                            else:
                                data["around"] = sentences[index-1]+sentences[index]
                            datas.append(data)
                            
        return datas

    '''文章分句处理'''
    def process_content(self, content):
        return [sentence for sentence in SentenceSplitter.split(content) if sentence]

    '''切分最小句'''
    def fined_sentence(self, sentence):
        return re.split(r'[？！，；]', sentence)



# 事件三元组
class SVOParser:
    def __init__(self):
        self.parser = DDParser(use_pos=True)
        print('loaded model')

    '''文章分句处理, 切分长句，冒号，分号，感叹号等做切分标识'''

    def split_sents(self, content):
        return [sentence for sentence in re.split(r'[？?！!。；;：:\n\r]', content) if sentence]

    '''句法分析---为句子中的每个词语维护一个保存句法依存儿子节点的字典'''
    def build_parse_child_dict(self, words, postags, rel_id, relation):
        child_dict_list = []
        format_parse_list = []
        for index in range(len(words)):
            child_dict = dict()
            for arc_index in range(len(rel_id)):
                if rel_id[arc_index] == index+1:   #arcs的索引从1开始
                    if rel_id[arc_index] in child_dict:
                        child_dict[relation[arc_index]].append(arc_index)
                    else:
                        child_dict[relation[arc_index]] = []
                        child_dict[relation[arc_index]].append(arc_index)
            child_dict_list.append(child_dict)
        heads = ['Root' if id == 0 else words[id - 1] for id in rel_id]  # 匹配依存父节点词语
        for i in range(len(words)):
            # ['ATT', '李克强', 0, 'nh', '总理', 1, 'n']
            a = [relation[i], words[i], i, postags[i], heads[i], rel_id[i]-1, postags[rel_id[i]-1]]
            format_parse_list.append(a)

        return child_dict_list, format_parse_list

    '''parser主函数'''
    def parser_main(self, sentence):
        res = self.parser.parse(sentence, )[0]
        words = res["word"]
        postags = res["postag"]
        rel_id = res["head"]
        relation = res["deprel"]

        child_dict_list, format_parse_list = self.build_parse_child_dict(words, postags, rel_id, relation)
        return words, postags, child_dict_list, format_parse_list

    """将所有的ATT进行合并"""
    def merge_ATT(self, words, postags, format_parse_list):
        words_ = words
        retain_nodes = set()
        ATTs = []
        ATT = []
        format_parse_list_ = []
        for parse in format_parse_list:
            dep = parse[0]
            if dep in ['ATT', 'ADV']:
                ATT += [parse[2], parse[5]]
            else:
                if ATT:
                    body = ''.join([words[i] for i in sorted(set(ATT))])
                    ATTs.append(body)
                    retain_nodes.add(sorted(set(ATT))[-1])
                    words_[sorted(set(ATT))[-1]] = body
                else:
                    retain_nodes.add(parse[2])
                ATT = []
        for indx, parse in enumerate(format_parse_list):
            if indx in retain_nodes:
                parse_ = [parse[0], words_[indx], indx, postags[indx], words_[parse[5]], parse[5], postags[parse[5]]]
                format_parse_list_.append(parse_)
        return words_, postags, format_parse_list_, retain_nodes

    """基于该结果，提取三元组"""
    def extract(self, words, postags, child_dict_list, arcs, retain_nodes):
        svos = []
        for index in range(len(postags)):
            if index not in retain_nodes:
                continue
            tmp = 1
            # 如果语义角色标记为空，则使用依存句法进行抽取
            if postags[index]:
                # 抽取以谓词为中心的事实三元组
                child_dict = child_dict_list[index]
                # 主谓宾
                if 'SBV' in child_dict and 'VOB' in child_dict:
                    # e1s = self.expand_e(words, postags, child_dict_list, child_dict['SBV'][0])
                    # e2s = self.expand_e(words, postags, child_dict_list, child_dict['VOB'][0])
                    r = words[index]
                    e1 = words[child_dict['SBV'][0]]
                    e2 = words[child_dict['VOB'][0]]
                    if e1.replace(' ', '') and e2.replace(' ', ''):
                        svos.append([e1, r, e2])

                # 含有介宾关系的主谓动补关系
                if 'SBV' in child_dict and 'CMP' in child_dict:
                    e1 = words[child_dict['SBV'][0]]
                    cmp_index = child_dict['CMP'][0]
                    r = words[index] + words[cmp_index]
                    if 'POB' in child_dict_list[cmp_index]:
                        e2 = words[child_dict_list[cmp_index]['POB'][0]]
                        if e1.replace(' ', '') and e2.replace(' ', ''):
                            svos.append([e1, r, e2])

        return svos

    '''三元组抽取主函数'''

    def ruler2(self, words, postags, child_dict_list, arcs):
        svos = []
        for index in range(len(postags)):
            tmp = 1
            # 先借助语义角色标注的结果，进行三元组抽取
            if tmp == 1:
                # 如果语义角色标记为空，则使用依存句法进行抽取
                # if postags[index] == 'v':
                if postags[index]:
                    # 抽取以谓词为中心的事实三元组
                    child_dict = child_dict_list[index]
                    # 主谓宾
                    if 'SBV' in child_dict and 'VOB' in child_dict:
                        r = words[index]
                        e1 = self.complete_e(words, postags, child_dict_list, child_dict['SBV'][0])
                        e2 = self.complete_e(words, postags, child_dict_list, child_dict['VOB'][0])
                        if e1.replace(' ', '') and e2.replace(' ', ''):
                            svos.append([e1, r, e2])

                    # 定语后置，动宾关系
                    relation = arcs[index][0]
                    head = arcs[index][2]
                    if relation == 'ATT':
                        if 'VOB' in child_dict:
                            e1 = self.complete_e(words, postags, child_dict_list, head - 1)
                            r = words[index]
                            e2 = self.complete_e(words, postags, child_dict_list, child_dict['VOB'][0])
                            temp_string = r + e2
                            if temp_string == e1[:len(temp_string)]:
                                e1 = e1[len(temp_string):]
                            if temp_string not in e1:
                                if e1.replace(' ', '') and e2.replace(' ', ''):
                                    svos.append([e1, r, e2])

                    # 含有介宾关系的主谓动补关系
                    if 'SBV' in child_dict and 'CMP' in child_dict:
                        e1 = self.complete_e(words, postags, child_dict_list, child_dict['SBV'][0])
                        cmp_index = child_dict['CMP'][0]
                        r = words[index] + words[cmp_index]
                        if 'POB' in child_dict_list[cmp_index]:
                            e2 = self.complete_e(words, postags, child_dict_list, child_dict_list[cmp_index]['POB'][0])
                            if e1.replace(' ', '') and e2.replace(' ', ''):
                                svos.append([e1, r, e2])
        return svos

    '''对找出的主语或者宾语进行扩展'''

    def complete_e(self, words, postags, child_dict_list, word_index):
        child_dict = child_dict_list[word_index]
        prefix = ''
        if 'ATT' in child_dict:
            for i in range(len(child_dict['ATT'])):
                prefix += self.complete_e(words, postags, child_dict_list, child_dict['ATT'][i])
        postfix = ''
        if postags[word_index] == 'v':
            if 'VOB' in child_dict:
                postfix += self.complete_e(words, postags, child_dict_list, child_dict['VOB'][0])
            if 'SBV' in child_dict:
                prefix = self.complete_e(words, postags, child_dict_list, child_dict['SBV'][0]) + prefix

        return prefix + words[word_index] + postfix

    '''程序主控函数'''

    def triples_main(self, content):
        sentences = self.split_sents(content)
        svos = []
        for sentence in sentences:
            # print(sentence)
            words, postags, child_dict_list, arcs = self.parser_main(sentence)
            svo = self.ruler2(words, postags, child_dict_list, arcs)
            svos += svo

        return svos

def write2excel(i,f_pass,f_list):
    # workbook = xlwt.Workbook(encoding = 'utf-8')
    # worksheet = workbook.add_sheet('sheet')

    if os.path.splitext(i)[1] == '.docx':
        doc = Document(f_pass+i)
        string = ""
        for para in doc.paragraphs :
            string = string + para.text
        
        print("loading model...")
        extractor1 = CausalityExractor()

        extractor2 = SVOParser()
        
        datas = extractor1.extract_main(string)
        # print(around)
        for data in datas:
            
            print('******'*4)
            print('cause', ''.join([word.split('/')[0] for word in data['cause'].split(' ') if word.split('/')[0]]))
            print('tag', data['tag'])
            print('effect', ''.join([word.split('/')[0] for word in data['effect'].split(' ') if word.split('/')[0]]))
            print("around:",data["around"])

            content_cause = "".join([word.split('/')[0] for word in data['cause'].split(' ') if word.split('/')[0]])
            content_effect = "".join([word.split('/')[0] for word in data['effect'].split(' ') if word.split('/')[0]])
            sentence_around = data["around"]
            k = f_list.index(i)+1

            print(k)
            # print(content_cause)
            # print(content_effect)


            # print(content)
            
            # if len(content_cause) < 100 and len(content_cause) > 5 and len(content_effect)<100 and len(content_effect) > 5:
            if True:
                
                
                worksheet.write(7*k,2, content_effect)
                worksheet.write(7*k,0, content_cause)
                worksheet.write(7*k,4, sentence_around)
                
                svos1 = extractor2.triples_main(content_cause)
                # print(svos1)

                k1,k2 = 7*k

                for svo1 in svos1:
                    worksheet.write(k2,1, str(svo1))
                    k2+=1
                    print("cause:")
                    print(svo1)
            
                svos2 = extractor2.triples_main(content_effect)
                
                for svo2 in svos2:
                    worksheet.write(k1,3, str(svo2))
                    k1+=1
                    print("effect:")
                    print(svo2)
                    
            # workbook.save('EventTriplesExtraction.xls')
    



if __name__ == '__main__':

    filePath = '../1_3diff/'
    f_list = os.listdir(filePath)

    p = Pool(5)
    f_pass = "../1_3diff/"

    workbook = xlwt.Workbook(encoding = 'utf-8')
    worksheet = workbook.add_sheet('sheet')
    worksheet.write(0,0, 'cause sentence')
    worksheet.write(0,1, 'cause triples')
    worksheet.write(0,2, 'effect sentence')
    worksheet.write(0,3, 'effect triples')
    worksheet.write(0,4, 'contex around')
    
    workbook.save('EventTriplesExtraction.xls')
    

    # for i in f_list[500:1000]:
    #     write2excel(i,f_pass,f_list)
    # #     p.apply_async(write2excel,args = (i,f_pass,f_list))
    # # p.close()
    # # p.join()

    # workbook.save('EventTriplesExtraction.xls')


    
    k = 1

    for i in f_list[500:1000]:
        # os.path.splitext():分离文件名与扩展名
        if os.path.splitext(i)[1] == '.docx':
            doc = Document("../1_3diff/"+i)
            string = ""
            for para in doc.paragraphs :
                string = string + para.text
            
            print("loading model...")
            extractor1 = CausalityExractor()

            extractor2 = SVOParser()
            
            datas = extractor1.extract_main(string)
            # print(around)
            for data in datas:
                
                print('******'*4)
                print('cause', ''.join([word.split('/')[0] for word in data['cause'].split(' ') if word.split('/')[0]]))
                print('tag', data['tag'])
                print('effect', ''.join([word.split('/')[0] for word in data['effect'].split(' ') if word.split('/')[0]]))
                print("around:",data["around"])

                print(f_list.index(i)+1)
              

                content_cause = "".join([word.split('/')[0] for word in data['cause'].split(' ') if word.split('/')[0]])
                content_effect = "".join([word.split('/')[0] for word in data['effect'].split(' ') if word.split('/')[0]])
                sentence_around = data["around"]
                # print(content)
                
                if len(content_cause) < 100 and len(content_cause) > 5 and len(content_effect)<100 and len(content_effect) > 5:
                # if True:
                    
                    worksheet.write(k,2, content_effect)
                    worksheet.write(k,0, content_cause)
                    worksheet.write(k,4, sentence_around)

                    
                    svos1 = extractor2.triples_main(content_cause)

                    k1 = k

                    for svo1 in svos1:
                        worksheet.write(k,1, str(svo1))
                        k+=1
                        print("cause:")
                        print(svo1)
                
                    svos2 = extractor2.triples_main(content_effect)
                    
                    for svo2 in svos2:
                        worksheet.write(k1,3, str(svo2))
                        k1+=1
                        print("effect:")
                        print(svo2)

                    if k1 > k:
                        k = k1
                    k += 1
            workbook.save('EventTriples7.xls')
    
    

